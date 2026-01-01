import os
import re
import logging
import zipfile
import tarfile
import tempfile
import mimetypes
from typing import List, Dict, Tuple, Set, Optional
from pathlib import Path
import chardet
import json
import csv

from config import IGNORED_PATTERNS, BLACKLISTED_DOMAINS
from link_utils import (
    normalize_url, is_url_ignored,
    analyze_link, extract_urls_from_text
)

# ======================
# Logging
# ======================

logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# ======================
# Constants
# ======================

# الأنواع المدعومة للملفات النصية
TEXT_FILE_EXTENSIONS = {
    '.txt', '.log', '.csv', '.json', '.xml', '.yaml', '.yml',
    '.html', '.htm', '.md', '.rst', '.ini', '.cfg', '.conf',
    '.php', '.js', '.py', '.java', '.cpp', '.c', '.h', '.cs',
    '.sql', '.sh', '.bash', '.ps1', '.bat', '.cmd'
}

# الأنواع المدعومة للملفات المضغوطة
ARCHIVE_FILE_EXTENSIONS = {
    '.zip', '.rar', '.7z', '.tar', '.gz', '.bz2', '.xz',
    '.tgz', '.tbz2', '.txz'
}

# الأنواع المدعومة للملفات المكتبية
OFFICE_FILE_EXTENSIONS = {
    '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx',
    '.odt', '.ods', '.odp', '.pdf'
}

# الأنواع المدعومة للملفات الأخرى
OTHER_SUPPORTED_EXTENSIONS = {
    '.sqlite', '.db', '.sqlite3', '.db3'
}

# جميع الأنواع المدعومة
ALL_SUPPORTED_EXTENSIONS = (
    TEXT_FILE_EXTENSIONS |
    ARCHIVE_FILE_EXTENSIONS |
    OFFICE_FILE_EXTENSIONS |
    OTHER_SUPPORTED_EXTENSIONS
)

# الأنماط الشائعة للعثور على الروابط في النص
URL_PATTERNS = [
    # روابط تيليجرام
    r't\.me/[A-Za-z0-9_+-]+',
    r'telegram\.me/[A-Za-z0-9_+-]+',
    r'tg://[A-Za-z0-9_?=&-]+',
    
    # روابط واتساب
    r'chat\.whatsapp\.com/[A-Za-z0-9_-]+',
    r'whatsapp\.com/channel/[A-Za-z0-9_-]+',
    r'wa\.me/[0-9]+',
    
    # روابط عامة
    r'https?://[^\s<>"\'{}|\\^`\[\]]+',
    r'www\.[^\s<>"\'{}|\\^`\[\]]+\.[^\s<>"\'{}|\\^`\[\]]+',
]

# ======================
# Text File Processing
# ======================

def detect_file_encoding(filepath: str) -> str:
    """كشف ترميز الملف"""
    try:
        with open(filepath, 'rb') as f:
            raw_data = f.read(10000)  # قراءة أول 10KB للكشف
            result = chardet.detect(raw_data)
            
            encoding = result.get('encoding', 'utf-8')
            confidence = result.get('confidence', 0)
            
            if confidence < 0.7:
                logger.warning(f"Low encoding confidence for {filepath}: {encoding} ({confidence})")
                return 'utf-8'
            
            # تحويل بعض الترميزات المشتركة
            if encoding.lower() in ['windows-1256', 'iso-8859-6', 'arabic']:
                return 'cp1256'
            
            return encoding.lower()
            
    except Exception as e:
        logger.error(f"Error detecting encoding for {filepath}: {e}")
        return 'utf-8'

def read_text_file(filepath: str, encoding: str = 'utf-8') -> str:
    """قراءة ملف نصي"""
    try:
        # محاولة قراءة بالترميز المحدد
        try:
            with open(filepath, 'r', encoding=encoding, errors='ignore') as f:
                return f.read()
        except UnicodeDecodeError:
            # محاولة ترميزات أخرى شائعة
            encodings_to_try = ['utf-8-sig', 'cp1256', 'latin-1', 'iso-8859-1']
            
            for enc in encodings_to_try:
                try:
                    with open(filepath, 'r', encoding=enc, errors='ignore') as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # إذا فشلت كل المحاولات، استخدم binary mode
            with open(filepath, 'rb') as f:
                content = f.read()
                return content.decode('utf-8', errors='ignore')
                
    except Exception as e:
        logger.error(f"Error reading file {filepath}: {e}")
        return ""

def extract_urls_from_text_content(text: str) -> List[str]:
    """استخراج الروابط من محتوى نصي"""
    if not text:
        return []
    
    urls = []
    
    # استخدام regex للعثور على جميع الروابط المحتملة
    for pattern in URL_PATTERNS:
        matches = re.findall(pattern, text, re.IGNORECASE)
        urls.extend(matches)
    
    # تنظيف الروابط وإضافة https:// إذا لزم الأمر
    cleaned_urls = []
    for url in urls:
        url = url.strip()
        
        # إضافة البروتوكول إذا لم يكن موجوداً
        if not url.startswith(('http://', 'https://')):
            if url.startswith('www.'):
                url = 'https://' + url
            elif '://' not in url:
                # روابط تيليجرام بدون بروتوكول
                if url.startswith(('t.me/', 'telegram.me/', 'tg://')):
                    url = 'https://' + url if not url.startswith('tg://') else url
        
        cleaned_urls.append(url)
    
    return list(set(cleaned_urls))  # إزالة التكرار

def process_text_file(filepath: str) -> Dict:
    """معالجة ملف نصي واستخراج الروابط"""
    results = {
        'filepath': filepath,
        'filename': os.path.basename(filepath),
        'file_size': os.path.getsize(filepath) if os.path.exists(filepath) else 0,
        'encoding': 'unknown',
        'total_urls_found': 0,
        'valid_urls': 0,
        'telegram_urls': 0,
        'whatsapp_urls': 0,
        'other_urls': 0,
        'ignored_urls': 0,
        'extracted_urls': [],
        'valid_urls_list': [],
        'error': None
    }
    
    try:
        # كشف ترميز الملف
        encoding = detect_file_encoding(filepath)
        results['encoding'] = encoding
        
        # قراءة الملف
        content = read_text_file(filepath, encoding)
        
        if not content:
            results['error'] = 'File is empty or cannot be read'
            return results
        
        # استخراج الروابط
        raw_urls = extract_urls_from_text_content(content)
        results['total_urls_found'] = len(raw_urls)
        
        # تحليل كل رابط
        for url in raw_urls:
            try:
                # تطبيع الرابط
                normalized = normalize_url(url)
                
                # التحقق من التجاهل
                ignored, ignore_reason = is_url_ignored(normalized)
                if ignored:
                    results['ignored_urls'] += 1
                    continue
                
                results['extracted_urls'].append(normalized)
                
                # تحليل الرابط
                analysis = analyze_link(normalized)
                
                if analysis['is_valid']:
                    results['valid_urls'] += 1
                    results['valid_urls_list'].append({
                        'url': normalized,
                        'platform': analysis['platform'],
                        'link_type': analysis['link_type'],
                        'should_collect': analysis['should_collect']
                    })
                    
                    # تعداد حسب المنصة
                    if analysis['platform'] == 'telegram':
                        results['telegram_urls'] += 1
                    elif analysis['platform'] == 'whatsapp':
                        results['whatsapp_urls'] += 1
                    else:
                        results['other_urls'] += 1
                
            except Exception as e:
                logger.error(f"Error processing URL {url}: {e}")
                continue
        
        logger.info(f"Processed text file: {filepath} - Found {results['valid_urls']} valid URLs")
        
    except Exception as e:
        logger.error(f"Error processing text file {filepath}: {e}")
        results['error'] = str(e)
    
    return results

# ======================
# Archive File Processing
# ======================

def extract_archive(filepath: str, extract_to: str) -> List[str]:
    """استخراج ملف مضغوط"""
    extracted_files = []
    
    try:
        if not os.path.exists(filepath):
            logger.error(f"Archive file not found: {filepath}")
            return extracted_files
        
        # إنشاء مجلد الاستخراج إذا لم يكن موجوداً
        os.makedirs(extract_to, exist_ok=True)
        
        # استخراج حسب نوع الملف
        if filepath.endswith('.zip'):
            with zipfile.ZipFile(filepath, 'r') as zip_ref:
                zip_ref.extractall(extract_to)
                extracted_files = zip_ref.namelist()
        
        elif filepath.endswith(('.tar', '.tar.gz', '.tgz', '.tar.bz2', '.tbz2', '.tar.xz', '.txz')):
            with tarfile.open(filepath, 'r:*') as tar_ref:
                tar_ref.extractall(extract_to)
                extracted_files = tar_ref.getnames()
        
        elif filepath.endswith('.rar'):
            try:
                import rarfile
                with rarfile.RarFile(filepath) as rar_ref:
                    rar_ref.extractall(extract_to)
                    extracted_files = rar_ref.namelist()
            except ImportError:
                logger.error("rarfile library not installed. Cannot extract RAR files.")
                return []
        
        elif filepath.endswith('.7z'):
            try:
                import py7zr
                with py7zr.SevenZipFile(filepath, 'r') as sz_ref:
                    sz_ref.extractall(extract_to)
                    extracted_files = sz_ref.getnames()
            except ImportError:
                logger.error("py7zr library not installed. Cannot extract 7z files.")
                return []
        
        logger.info(f"Extracted {len(extracted_files)} files from {filepath}")
        
    except Exception as e:
        logger.error(f"Error extracting archive {filepath}: {e}")
    
    return extracted_files

def process_archive_file(filepath: str) -> Dict:
    """معالجة ملف مضغوط واستخراج الروابط من محتوياته"""
    results = {
        'filepath': filepath,
        'filename': os.path.basename(filepath),
        'file_size': os.path.getsize(filepath) if os.path.exists(filepath) else 0,
        'extracted_files': 0,
        'processed_files': 0,
        'total_urls_found': 0,
        'valid_urls': 0,
        'telegram_urls': 0,
        'whatsapp_urls': 0,
        'other_urls': 0,
        'ignored_urls': 0,
        'valid_urls_list': [],
        'extracted_files_info': [],
        'error': None
    }
    
    try:
        # إنشاء مجلد مؤقت للاستخراج
        with tempfile.TemporaryDirectory() as temp_dir:
            # استخراج الملف المضغوط
            extracted_files = extract_archive(filepath, temp_dir)
            results['extracted_files'] = len(extracted_files)
            
            if not extracted_files:
                results['error'] = 'No files extracted or archive is empty'
                return results
            
            # معالجة كل ملف مستخرج
            for extracted_file in extracted_files:
                extracted_path = os.path.join(temp_dir, extracted_file)
                
                # التأكد من أن المسار موجود وليس مجلداً
                if not os.path.isfile(extracted_path):
                    continue
                
                # التحقق من نوع الملف
                file_ext = os.path.splitext(extracted_file)[1].lower()
                
                file_info = {
                    'filename': extracted_file,
                    'filepath': extracted_path,
                    'file_size': os.path.getsize(extracted_path),
                    'processed': False,
                    'urls_found': 0,
                    'valid_urls': 0,
                    'error': None
                }
                
                try:
                    # معالجة الملفات النصية فقط
                    if file_ext in TEXT_FILE_EXTENSIONS:
                        # معالجة الملف النصي
                        text_results = process_text_file(extracted_path)
                        
                        file_info['processed'] = True
                        file_info['urls_found'] = text_results['total_urls_found']
                        file_info['valid_urls'] = text_results['valid_urls']
                        
                        # تحديث النتائج الإجمالية
                        results['processed_files'] += 1
                        results['total_urls_found'] += text_results['total_urls_found']
                        results['valid_urls'] += text_results['valid_urls']
                        results['telegram_urls'] += text_results['telegram_urls']
                        results['whatsapp_urls'] += text_results['whatsapp_urls']
                        results['other_urls'] += text_results['other_urls']
                        results['ignored_urls'] += text_results['ignored_urls']
                        
                        # إضافة الروابط الصالحة
                        for url_info in text_results.get('valid_urls_list', []):
                            if url_info['should_collect']:
                                results['valid_urls_list'].append({
                                    'url': url_info['url'],
                                    'source_file': extracted_file,
                                    'platform': url_info['platform'],
                                    'link_type': url_info['link_type']
                                })
                    
                    else:
                        file_info['error'] = f'Unsupported file type: {file_ext}'
                
                except Exception as e:
                    file_info['error'] = str(e)
                    logger.error(f"Error processing extracted file {extracted_file}: {e}")
                
                results['extracted_files_info'].append(file_info)
        
        logger.info(f"Processed archive: {filepath} - Found {results['valid_urls']} valid URLs in {results['processed_files']} files")
        
    except Exception as e:
        logger.error(f"Error processing archive file {filepath}: {e}")
        results['error'] = str(e)
    
    return results

# ======================
# Office File Processing
# ======================

def process_office_file(filepath: str) -> Dict:
    """معالجة ملف مكتبي (Word, Excel, PDF, etc.)"""
    results = {
        'filepath': filepath,
        'filename': os.path.basename(filepath),
        'file_size': os.path.getsize(filepath) if os.path.exists(filepath) else 0,
        'total_urls_found': 0,
        'valid_urls': 0,
        'telegram_urls': 0,
        'whatsapp_urls': 0,
        'other_urls': 0,
        'ignored_urls': 0,
        'valid_urls_list': [],
        'error': None
    }
    
    try:
        content = ""
        file_ext = os.path.splitext(filepath)[1].lower()
        
        # معالجة حسب نوع الملف
        if file_ext == '.pdf':
            try:
                import PyPDF2
                with open(filepath, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        content += page.extract_text() or ""
            except ImportError:
                logger.warning("PyPDF2 not installed. Cannot extract text from PDF.")
                results['error'] = 'PDF extraction requires PyPDF2 library'
                return results
            except Exception as e:
                logger.error(f"Error extracting text from PDF: {e}")
                results['error'] = str(e)
                return results
        
        elif file_ext in ['.docx', '.pptx', '.xlsx']:
            try:
                from docx import Document
                doc = Document(filepath)
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
                
                # معالجة الجداول
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            content += cell.text + " "
            except ImportError:
                logger.warning("python-docx not installed. Cannot extract text from Office files.")
                results['error'] = 'Office file extraction requires python-docx library'
                return results
            except Exception as e:
                logger.error(f"Error extracting text from Office file: {e}")
                results['error'] = str(e)
                return results
        
        elif file_ext in ['.doc', '.ppt', '.xls']:
            # الملفات القديمة تحتاج إلى تحويل
            results['error'] = 'Old Office format (.doc, .ppt, .xls) requires conversion'
            return results
        
        else:
            results['error'] = f'Unsupported Office file type: {file_ext}'
            return results
        
        if not content:
            results['error'] = 'No text content extracted'
            return results
        
        # استخراج الروابط من المحتوى
        urls = extract_urls_from_text_content(content)
        results['total_urls_found'] = len(urls)
        
        # تحليل كل رابط
        for url in urls:
            try:
                # تطبيع الرابط
                normalized = normalize_url(url)
                
                # التحقق من التجاهل
                ignored, ignore_reason = is_url_ignored(normalized)
                if ignored:
                    results['ignored_urls'] += 1
                    continue
                
                # تحليل الرابط
                analysis = analyze_link(normalized)
                
                if analysis['is_valid'] and analysis['should_collect']:
                    results['valid_urls'] += 1
                    results['valid_urls_list'].append({
                        'url': normalized,
                        'platform': analysis['platform'],
                        'link_type': analysis['link_type']
                    })
                    
                    # تعداد حسب المنصة
                    if analysis['platform'] == 'telegram':
                        results['telegram_urls'] += 1
                    elif analysis['platform'] == 'whatsapp':
                        results['whatsapp_urls'] += 1
                    else:
                        results['other_urls'] += 1
                
            except Exception as e:
                logger.error(f"Error processing URL {url}: {e}")
                continue
        
        logger.info(f"Processed office file: {filepath} - Found {results['valid_urls']} valid URLs")
        
    except Exception as e:
        logger.error(f"Error processing office file {filepath}: {e}")
        results['error'] = str(e)
    
    return results

# ======================
# Database File Processing
# ======================

def process_sqlite_file(filepath: str) -> Dict:
    """معالجة ملف SQLite واستخراج الروابط"""
    results = {
        'filepath': filepath,
        'filename': os.path.basename(filepath),
        'file_size': os.path.getsize(filepath) if os.path.exists(filepath) else 0,
        'tables_found': 0,
        'total_urls_found': 0,
        'valid_urls': 0,
        'telegram_urls': 0,
        'whatsapp_urls': 0,
        'other_urls': 0,
        'ignored_urls': 0,
        'valid_urls_list': [],
        'error': None
    }
    
    try:
        import sqlite3
        
        if not os.path.exists(filepath):
            results['error'] = 'File not found'
            return results
        
        # الاتصال بقاعدة البيانات
        conn = sqlite3.connect(filepath)
        cursor = conn.cursor()
        
        # الحصول على قائمة الجداول
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        tables = cursor.fetchall()
        results['tables_found'] = len(tables)
        
        # البحث في كل جدول عن الروابط
        urls_found = set()
        
        for table in tables:
            table_name = table[0]
            
            try:
                # الحصول على أسماء الأعمدة
                cursor.execute(f"PRAGMA table_info({table_name})")
                columns = cursor.fetchall()
                
                # البحث عن أعمدة قد تحتوي على روابط
                url_columns = []
                for col in columns:
                    col_name = col[1].lower()
                    if any(keyword in col_name for keyword in ['url', 'link', 'href', 'telegram', 'whatsapp', 't.me']):
                        url_columns.append(col[1])
                
                if url_columns:
                    # استعلام البيانات من الأعمدة ذات الصلة
                    for col in url_columns:
                        cursor.execute(f"SELECT DISTINCT {col} FROM {table_name} WHERE {col} IS NOT NULL AND {col} != ''")
                        rows = cursor.fetchall()
                        
                        for row in rows:
                            value = str(row[0])
                            # استخراج الروابط من النص
                            extracted_urls = extract_urls_from_text_content(value)
                            urls_found.update(extracted_urls)
                
                # البحث في جميع الأعمدة النصية
                for col in columns:
                    col_name = col[1]
                    col_type = col[2].lower()
                    
                    if 'text' in col_type or 'char' in col_type or 'varchar' in col_type:
                        try:
                            cursor.execute(f"SELECT {col_name} FROM {table_name} WHERE {col_name} LIKE '%t.me%' OR {col_name} LIKE '%telegram.me%' OR {col_name} LIKE '%whatsapp.com%' LIMIT 100")
                            rows = cursor.fetchall()
                            
                            for row in rows:
                                value = str(row[0])
                                extracted_urls = extract_urls_from_text_content(value)
                                urls_found.update(extracted_urls)
                        except:
                            continue
            
            except Exception as e:
                logger.error(f"Error processing table {table_name}: {e}")
                continue
        
        conn.close()
        
        results['total_urls_found'] = len(urls_found)
        
        # تحليل كل رابط
        for url in urls_found:
            try:
                # تطبيع الرابط
                normalized = normalize_url(url)
                
                # التحقق من التجاهل
                ignored, ignore_reason = is_url_ignored(normalized)
                if ignored:
                    results['ignored_urls'] += 1
                    continue
                
                # تحليل الرابط
                analysis = analyze_link(normalized)
                
                if analysis['is_valid'] and analysis['should_collect']:
                    results['valid_urls'] += 1
                    results['valid_urls_list'].append({
                        'url': normalized,
                        'platform': analysis['platform'],
                        'link_type': analysis['link_type']
                    })
                    
                    # تعداد حسب المنصة
                    if analysis['platform'] == 'telegram':
                        results['telegram_urls'] += 1
                    elif analysis['platform'] == 'whatsapp':
                        results['whatsapp_urls'] += 1
                    else:
                        results['other_urls'] += 1
                
            except Exception as e:
                logger.error(f"Error processing URL {url}: {e}")
                continue
        
        logger.info(f"Processed SQLite file: {filepath} - Found {results['valid_urls']} valid URLs in {results['tables_found']} tables")
        
    except ImportError:
        results['error'] = 'SQLite processing requires sqlite3 module'
    except Exception as e:
        logger.error(f"Error processing SQLite file {filepath}: {e}")
        results['error'] = str(e)
    
    return results

# ======================
# Generic File Processing
# ======================

def get_file_type(filepath: str) -> str:
    """الحصول على نوع الملف"""
    if not os.path.exists(filepath):
        return 'not_found'
    
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext in TEXT_FILE_EXTENSIONS:
        return 'text'
    elif ext in ARCHIVE_FILE_EXTENSIONS:
        return 'archive'
    elif ext in OFFICE_FILE_EXTENSIONS:
        return 'office'
    elif ext in OTHER_SUPPORTED_EXTENSIONS:
        return 'database'
    else:
        return 'unsupported'

def process_file(filepath: str) -> Dict:
    """معالجة ملف واستخراج الروابط منه"""
    file_type = get_file_type(filepath)
    
    if not os.path.exists(filepath):
        return {
            'filepath': filepath,
            'filename': os.path.basename(filepath),
            'error': 'File not found',
            'file_type': 'not_found'
        }
    
    logger.info(f"Processing file: {filepath} (Type: {file_type})")
    
    if file_type == 'text':
        return process_text_file(filepath)
    
    elif file_type == 'archive':
        return process_archive_file(filepath)
    
    elif file_type == 'office':
        return process_office_file(filepath)
    
    elif file_type == 'database':
        return process_sqlite_file(filepath)
    
    else:
        return {
            'filepath': filepath,
            'filename': os.path.basename(filepath),
            'file_size': os.path.getsize(filepath),
            'error': f'Unsupported file type: {os.path.splitext(filepath)[1]}',
            'file_type': 'unsupported'
        }

def process_directory(directory_path: str, recursive: bool = True) -> Dict:
    """معالجة مجلد كامل واستخراج الروابط من جميع الملفات"""
    results = {
        'directory': directory_path,
        'total_files': 0,
        'processed_files': 0,
        'unsupported_files': 0,
        'failed_files': 0,
        'total_urls_found': 0,
        'valid_urls': 0,
        'telegram_urls': 0,
        'whatsapp_urls': 0,
        'other_urls': 0,
        'files_processed': [],
        'valid_urls_list': [],
        'summary_by_type': {
            'text': 0,
            'archive': 0,
            'office': 0,
            'database': 0,
            'unsupported': 0
        }
    }
    
    try:
        if not os.path.exists(directory_path) or not os.path.isdir(directory_path):
            results['error'] = 'Directory not found or not a directory'
            return results
        
        # جمع جميع الملفات
        all_files = []
        
        if recursive:
            for root, dirs, files in os.walk(directory_path):
                for file in files:
                    all_files.append(os.path.join(root, file))
        else:
            for item in os.listdir(directory_path):
                item_path = os.path.join(directory_path, item)
                if os.path.isfile(item_path):
                    all_files.append(item_path)
        
        results['total_files'] = len(all_files)
        
        # معالجة كل ملف
        for filepath in all_files:
            try:
                file_type = get_file_type(filepath)
                results['summary_by_type'][file_type] += 1
                
                # تجاهل الملفات غير المدعومة
                if file_type == 'unsupported':
                    results['unsupported_files'] += 1
                    continue
                
                # معالجة الملف
                file_results = process_file(filepath)
                
                file_summary = {
                    'filepath': filepath,
                    'filename': os.path.basename(filepath),
                    'file_type': file_type,
                    'file_size': file_results.get('file_size', 0),
                    'urls_found': file_results.get('total_urls_found', 0),
                    'valid_urls': file_results.get('valid_urls', 0),
                    'error': file_results.get('error')
                }
                
                results['files_processed'].append(file_summary)
                results['processed_files'] += 1
                
                if not file_results.get('error'):
                    results['total_urls_found'] += file_results.get('total_urls_found', 0)
                    results['valid_urls'] += file_results.get('valid_urls', 0)
                    results['telegram_urls'] += file_results.get('telegram_urls', 0)
                    results['whatsapp_urls'] += file_results.get('whatsapp_urls', 0)
                    results['other_urls'] += file_results.get('other_urls', 0)
                    
                    # جمع الروابط الصالحة
                    for url_info in file_results.get('valid_urls_list', []):
                        results['valid_urls_list'].append({
                            'url': url_info['url'],
                            'source_file': os.path.basename(filepath),
                            'platform': url_info.get('platform'),
                            'link_type': url_info.get('link_type')
                        })
                else:
                    results['failed_files'] += 1
                
                logger.info(f"Processed: {filepath} - {file_results.get('valid_urls', 0)} valid URLs")
                
            except Exception as e:
                logger.error(f"Error processing file {filepath}: {e}")
                results['failed_files'] += 1
                continue
        
        logger.info(f"Directory processing completed: {results['valid_urls']} valid URLs found in {results['processed_files']} files")
        
    except Exception as e:
        logger.error(f"Error processing directory {directory_path}: {e}")
        results['error'] = str(e)
    
    return results

# ======================
# Export Functions
# ======================

def export_extracted_urls(results: Dict, output_file: str = None) -> str:
    """تصدير الروابط المستخرجة إلى ملف"""
    try:
        if not results.get('valid_urls_list'):
            logger.warning("No valid URLs to export")
            return ""
        
        # إنشاء اسم الملف إذا لم يتم توفيره
        if not output_file:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_file = f"extracted_urls_{timestamp}.txt"
        
        # تجميع الروابط حسب المنصة والنوع
        telegram_groups = []
        telegram_channels = []
        whatsapp_groups = []
        other_urls = []
        
        for url_info in results['valid_urls_list']:
            url = url_info['url']
            platform = url_info.get('platform', '')
            link_type = url_info.get('link_type', '')
            
            if platform == 'telegram':
                if link_type in ['public_group', 'private_group']:
                    telegram_groups.append(url)
                elif link_type == 'channel':
                    telegram_channels.append(url)
                else:
                    other_urls.append(url)
            elif platform == 'whatsapp' and link_type == 'group':
                whatsapp_groups.append(url)
            else:
                other_urls.append(url)
        
        # كتابة الملف
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"# Extracted URLs Report\n")
            f.write(f"# Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"# Total Valid URLs: {len(results['valid_urls_list'])}\n")
            f.write("#" * 60 + "\n\n")
            
            if telegram_groups:
                f.write(f"# Telegram Groups ({len(telegram_groups)})\n")
                f.write("#" * 40 + "\n")
                for url in telegram_groups:
                    f.write(f"{url}\n")
                f.write("\n")
            
            if telegram_channels:
                f.write(f"# Telegram Channels ({len(telegram_channels)})\n")
                f.write("#" * 40 + "\n")
                for url in telegram_channels:
                    f.write(f"{url}\n")
                f.write("\n")
            
            if whatsapp_groups:
                f.write(f"# WhatsApp Groups ({len(whatsapp_groups)})\n")
                f.write("#" * 40 + "\n")
                for url in whatsapp_groups:
                    f.write(f"{url}\n")
                f.write("\n")
            
            if other_urls:
                f.write(f"# Other URLs ({len(other_urls)})\n")
                f.write("#" * 40 + "\n")
                for url in other_urls:
                    f.write(f"{url}\n")
        
        logger.info(f"Exported {len(results['valid_urls_list'])} URLs to {output_file}")
        return output_file
        
    except Exception as e:
        logger.error(f"Error exporting URLs: {e}")
        return ""

def save_processing_report(results: Dict, report_file: str = None) -> str:
    """حفظ تقرير المعالجة إلى ملف JSON"""
    try:
        # إنشاء اسم الملف إذا لم يتم توفيره
        if not report_file:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            report_file = f"processing_report_{timestamp}.json"
        
        # تحضير البيانات للتصدير
        export_data = {
            'report_date': datetime.now().isoformat(),
            'summary': {
                'total_files': results.get('total_files', 0),
                'processed_files': results.get('processed_files', 0),
                'unsupported_files': results.get('unsupported_files', 0),
                'failed_files': results.get('failed_files', 0),
                'total_urls_found': results.get('total_urls_found', 0),
                'valid_urls': results.get('valid_urls', 0),
                'telegram_urls': results.get('telegram_urls', 0),
                'whatsapp_urls': results.get('whatsapp_urls', 0),
                'other_urls': results.get('other_urls', 0)
            },
            'file_types_summary': results.get('summary_by_type', {}),
            'files_processed': results.get('files_processed', []),
            'valid_urls_count': len(results.get('valid_urls_list', [])),
            'error': results.get('error')
        }
        
        # كتابة الملف
        with open(report_file, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved processing report to {report_file}")
        return report_file
        
    except Exception as e:
        logger.error(f"Error saving processing report: {e}")
        return ""

# ======================
# Test Functions
# ======================

def test_file_extraction():
    """اختبار استخراج الروابط من الملفات"""
    print("🔧 اختبار استخراج الروابط من الملفات...")
    print("=" * 80)
    
    # إنشاء ملف نصي تجريبي للاختبار
    test_content = """
    Telegram groups:
    https://t.me/test_group_1
    https://t.me/+invite123
    t.me/test_group_2
    
    WhatsApp groups:
    https://chat.whatsapp.com/group1
    https://whatsapp.com/channel/channel1
    
    Other links:
    https://t.me/channel_news (channel)
    https://facebook.com/groups/test
    https://discord.gg/test
    
    Invalid/ignored:
    https://t.me/c/1234567890 (private channel)
    https://t.me/botfather (bot)
    """
    
    # حفظ المحتوى في ملف مؤقت
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
        f.write(test_content)
        test_file = f.name
    
    try:
        print(f"\n📄 اختبار معالجة ملف نصي: {test_file}")
        
        # معالجة الملف
        results = process_text_file(test_file)
        
        print(f"\n📊 النتائج:")
        print(f"   • اسم الملف: {results['filename']}")
        print(f"   • حجم الملف: {results['file_size']} bytes")
        print(f"   • الترميز: {results['encoding']}")
        print(f"   • الروابط الموجودة: {results['total_urls_found']}")
        print(f"   • الروابط الصالحة: {results['valid_urls']}")
        print(f"   • مجموعات تيليجرام: {results['telegram_urls']}")
        print(f"   • مجموعات واتساب: {results['whatsapp_urls']}")
        print(f"   • روابط أخرى: {results['other_urls']}")
        print(f"   • الروابط المتجاهلة: {results['ignored_urls']}")
        
        if results.get('error'):
            print(f"   • الأخطاء: {results['error']}")
        
        print(f"\n🔗 الروابط الصالحة:")
        for url_info in results.get('valid_urls_list', []):
            print(f"   • {url_info['url']} ({url_info['platform']}/{url_info['link_type']})")
        
        # اختبار التصدير
        print(f"\n💾 اختبار التصدير...")
        export_file = export_extracted_urls(results, 'test_export.txt')
        if export_file and os.path.exists(export_file):
            print(f"   ✅ تم التصدير إلى: {export_file}")
            
            # قراءة الملف المصدر لعرض المحتوى
            with open(export_file, 'r', encoding='utf-8') as f:
                print(f"\n📄 محتوى الملف المصدر:")
                print("-" * 40)
                print(f.read()[:500])  # عرض أول 500 حرف فقط
                print("-" * 40)
            
            # تنظيف
            os.remove(export_file)
        
        print("\n✅ اختبار استخراج الملفات اكتمل بنجاح!")
        
    finally:
        # تنظيف الملف المؤقت
        if os.path.exists(test_file):
            os.remove(test_file)

# ======================
# Main Entry Point
# ======================

if __name__ == "__main__":
    import sys
    
    print("🚀 تشغيل أدوات استخراج الروابط من الملفات...")
    
    # اختبار الدوال الأساسية
    test_file_extraction()
    
    print("\n" + "=" * 80)
    print("📋 الأنواع المدعومة:")
    
    print("\n📄 الملفات النصية:")
    for ext in sorted(TEXT_FILE_EXTENSIONS):
        print(f"   • {ext}")
    
    print("\n📦 الملفات المضغوطة:")
    for ext in sorted(ARCHIVE_FILE_EXTENSIONS):
        print(f"   • {ext}")
    
    print("\n📊 الملفات المكتبية:")
    for ext in sorted(OFFICE_FILE_EXTENSIONS):
        print(f"   • {ext}")
    
    print("\n🗄️ ملفات قواعد البيانات:")
    for ext in sorted(OTHER_SUPPORTED_EXTENSIONS):
        print(f"   • {ext}")
    
    print(f"\n🎯 إجمالي الأنواع المدعومة: {len(ALL_SUPPORTED_EXTENSIONS)} نوع")
    
    print("\n✅ أدوات استخراج الروابط جاهزة للعمل!")
